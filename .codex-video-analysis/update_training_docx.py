from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell(cell, text, size=10.5, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_table_fixed_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def remove_table(table):
    table._element.getparent().remove(table._element)


def clear_rows_after(table, keep_count):
    for row in list(table.rows)[keep_count:]:
        table._tbl.remove(row._tr)


def clone_row(table, source_row):
    new_tr = deepcopy(source_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def main(src, out):
    doc = Document(src)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    # Fill short-video planning table.
    plan = doc.tables[0]
    set_table_fixed_width(plan, [1.2, 8.9])
    set_cell(plan.cell(0, 0), "短视频标题：3分钟上手 AI API 接入", bold=True)
    set_cell(plan.cell(0, 1), "短视频标题：3分钟上手 AI API 接入", bold=True)
    set_cell(
        plan.cell(1, 1),
        "本短视频围绕团队网站 APIusPro.cn 展开，主题是“快速找到并接入合适的 AI API”。创作背景是很多开发者在选择 AI API 时，需要同时确认官网、价格、免费额度、购买方式、接入步骤和常见问题，信息分散导致上手成本高。视频以 APIusPro.cn 为入口，展示网站如何按真实使用路径整理 API 选型、购买与接入教程，并用 DeepSeek API 的创建密钥和 CC Switch 配置流程作为案例，帮助观众理解从搜索到测试的完整路径。创作目的在于突出网站的实用性，引导用户访问 apiuspro.cn 搜索并测试所需 API。",
    )
    set_cell(
        plan.cell(2, 1),
        "视频定位为“产品介绍 + 操作演示 + 教程引导”的实用型短视频，面向需要快速接入 AI API 的开发者、独立开发者和 AI 工具使用者。创意上采用问题开场：“想在三分钟内找到并接入合适的 AI API 吗？”先抓住用户痛点，再用网站页面和真实操作流程证明解决方案。内容不做空泛宣传，而是通过 API 列表、价格与免费额度、购买教程、接入示例、本地部署指南等具体信息建立可信度。后半段以 DeepSeek 平台创建 API Key、在 CC Switch 中粘贴密钥并选择模型版本为实操重点，让观众看到完整接入路径。结尾补充小规模测试、额度限速、监控和密钥管理提醒，使视频既有转化引导，也符合安全使用规范。",
    )
    set_cell(
        plan.cell(3, 1),
        "视频结构分为四段。第一段为痛点与网站定位，用 APIusPro.cn 首页和 API 列表页面说明网站能整理 AI API 的选型、购买和接入教程。第二段为核心功能展示，依次突出 API 列表、官网入口、价格、免费额度、购买教程、接入示例、本地部署指南和常见问题。第三段为案例演示，以 DeepSeek API 为例，展示搜索模型、打开详情、查看免费额度、进入 DeepSeek 平台、查看使用额度、创建并复制 API Key 的过程。第四段为工具配置与行动号召，演示在 CC Switch 中添加 DeepSeek、粘贴 API Key、调整模型版本并完成设置，最后提醒用户先小规模测试，关注额度、限速和密钥安全，并引导访问 apiuspro.cn 搜索想用的 API。",
    )
    set_cell(
        plan.cell(4, 1),
        "画面以电脑录屏为主，搭配字幕重点提示和页面操作高亮。开头使用 APIusPro.cn 首页和 API 列表作为主画面，突出“按访问条件和接入难度选择 AI API”的信息层级。中段切换到 DeepSeek 教程页和 DeepSeek 控制台，画面重点放在用量信息、API keys、创建密钥等关键按钮上。后段展示 CC Switch 配置窗口，用近景突出选择 DeepSeek、粘贴 API Key、调整模型版本和一键设置等动作。剪辑节奏保持快速清晰，字幕采用口语化但规范的表达，删除过于随意的词语，将“卡密”统一为“API Key / 密钥”。背景音乐建议使用轻快科技感音乐，按钮点击处可配合轻微提示音，结尾保留网站域名和行动号召。",
    )

    # Rebuild storyboard table instead of fighting the original narrow template.
    remove_table(doc.tables[1])
    shots = [
        ("1", "全景", "0-4秒", "静态录屏+字幕进入", "打开 APIusPro.cn 首页，画面出现“想在三分钟内找到并接入合适的 AI API 吗？”", "轻快科技感音乐起", "网站首页居中，标题区域占画面主体"),
        ("2", "中景", "4-9秒", "鼠标引导", "展示 API 知识站定位：按真实使用路径整理 API 选型、购买与接入教程。", "旁白说明网站定位", "页面导航和列表区域同时入镜"),
        ("3", "近景", "9-14秒", "页面轻微滚动", "展示 API 列表，突出官网、价格和免费额度说明。", "点击音效", "表格信息放大，免费额度字段醒目"),
        ("4", "中景", "14-18秒", "横向切换", "展示购买教程入口，说明支持支付与发票相关信息。", "旁白承接", "教程入口位于画面中心"),
        ("5", "中近景", "18-23秒", "向下滚动", "展示接入示例、本地部署指南、示例代码和常见问题。", "音乐保持稳定", "代码和步骤区域作为视觉重点"),
        ("6", "近景", "23-31秒", "鼠标点击", "搜索 DeepSeek，打开详情页，查看免费额度和接入步骤。", "提示音+旁白", "搜索框、详情页标题、免费额度依次出现"),
        ("7", "近景", "31-40秒", "页面切换", "进入 DeepSeek 平台，查看使用额度，打开 API keys 并点击创建 API Key。", "按钮点击音效", "用量数字和创建按钮居中突出"),
        ("8", "特写", "40-46秒", "弹窗停留", "输入名称并创建密钥，复制生成的 API Key。字幕使用“密钥/API Key”，避免使用“卡密”。", "短提示音", "弹窗和复制按钮占画面中心"),
        ("9", "中景", "46-55秒", "窗口切换+点击", "打开 CC Switch，点击添加，选择 DeepSeek，粘贴刚复制的 API Key，并按需求选择模型版本。", "节奏略加快", "工具窗口居中，关键输入框放大"),
        ("10", "全景", "55-60秒", "回到网站+收尾字幕", "提醒先小规模测试，关注免费额度、限速、监控和密钥管理；结尾出现 apiuspro.cn 行动号召。", "音乐收束", "网站域名和搜索框作为最后一屏"),
    ]
    script = doc.add_table(rows=4 + len(shots), cols=7)
    script.style = "Table Grid"
    widths = [0.45, 0.65, 0.8, 0.95, 3.9, 1.25, 2.1]
    set_table_fixed_width(script, widths)
    script.cell(0, 0).merge(script.cell(0, 6))
    script.cell(1, 1).merge(script.cell(1, 6))
    script.cell(2, 1).merge(script.cell(2, 6))
    set_cell(script.cell(0, 0), "短视频脚本", bold=True)
    set_cell(script.cell(1, 0), "标题", bold=True)
    set_cell(script.cell(1, 1), "3分钟上手 AI API 接入", bold=True)
    set_cell(script.cell(2, 0), "视频时长", bold=True)
    set_cell(script.cell(2, 1), "约60秒以内（由原75秒素材压缩剪辑）", bold=True)
    for idx, header in enumerate(["镜号", "景别", "时长", "运镜方式", "画面内容", "音乐音效", "画面构图"]):
        set_cell(script.cell(3, idx), header, size=9, bold=True)
    for row_idx, shot in enumerate(shots, start=4):
        for idx, value in enumerate(shot):
            set_cell(script.cell(row_idx, idx), value, size=8)
            set_cell_width(script.cell(row_idx, idx), widths[idx])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=10.5, bold=run.bold)

    doc.save(out)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
